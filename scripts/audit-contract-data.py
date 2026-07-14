#!/usr/bin/env python3
"""
Contract data audit: compare docx source files against customer.json / supplier.json.

Reports missing articles, extra articles, and clause (項) count mismatches.

Usage:
    python3 scripts/audit-contract-data.py

Requirements:
    pip install python-docx

Docx files expected in data/contracts/:
    BP-IGUAZUビジネスパートナー契約書.docx  → customer.json / IGUAZUビジネスパートナー契約書（BP）
    BB-売買取引基本契約書.docx             → customer.json / 売買取引基本契約書（BB）
    KBKF-購買基本契約書.docx              → supplier.json / 購買基本契約書(業務委託条項含む)（KB/KF）
    KFKM-業務委託基本契約書.docx           → supplier.json / 業務委託基本契約書（KF/KM）

Notes:
    - KBKF has two sections: main contract (第1-27条) + ［業務委託条項］ (第1-32条 → 別紙-第N条)
    - KFKM has two sections: main contract (第1-43条) + 機密保持契約書 (第1-14条 → 添付1 機密保持契約書 第N条)
    - KFKM 添付2 誓約書 (1-15) is not clause-counted (each item is its own top-level entry)
    - BB uses mixed styles: （title） List Paragraph headers, LP clauses or Normal N.\t clauses
    - BB: single List Paragraph item = article body (stored in content, 0 clauses) — OK
"""

import json
import re
from pathlib import Path

try:
    from docx import Document
except ImportError:
    print("ERROR: python-docx not installed. Run: pip install python-docx")
    raise

DATA = Path(__file__).parent.parent / 'data' / 'contracts'


# ── Helpers ─────────────────────────────────────────────────────────────────

def to_half(t):
    return t.translate(str.maketrans('０１２３４５６７８９', '0123456789'))

ART_RE    = re.compile(r'^第(\d+)条')
CLAUSE_RE = re.compile(r'^(\d+)\.')
GONUM_RE  = re.compile(r'^[(（]\d+[)）]')   # (1) or （1） sub-items


# ── Parsers ──────────────────────────────────────────────────────────────────

def extract_articles_normal(paragraphs):
    """Parse BP/KBKF/KFKM style: 第N条 headers, N. clause lines.
    Returns {第N条: max_clause_num}.
    """
    articles, cur = {}, None
    for _style, text in paragraphs:
        if not text:
            continue
        m = ART_RE.match(to_half(text))
        if m:
            cur = f'第{int(m.group(1))}条'
            articles.setdefault(cur, 0)
            continue
        if cur is None:
            continue
        if GONUM_RE.match(text):
            continue
        m2 = CLAUSE_RE.match(text)
        if m2:
            n = int(m2.group(1))
            if 1 <= n <= 35:
                articles[cur] = max(articles[cur], n)
    return articles


def extract_articles_bb(path):
    """Parse BB style: （title）List Paragraph headers, LP items or Normal N.\\t lines.
    Single LP item = article body (not a clause). Returns {title: clause_count}.
    """
    articles, cur, lp_cnt = {}, None, 0

    for para in Document(path).paragraphs:
        s = para.style.name
        t = para.text.strip()
        raw = para.text
        if not t:
            continue

        # Article header: （title） in List Paragraph
        if s == 'List Paragraph' and t.startswith('（') and t.endswith('）'):
            cur = t[1:-1]
            lp_cnt = 0
            articles.setdefault(cur, 0)
            continue

        if cur is None:
            continue
        if s == '号':
            continue
        if GONUM_RE.match(t):
            continue

        if s == 'List Paragraph':
            lp_cnt += 1
            articles[cur] = lp_cnt
            continue

        if s == 'Normal':
            m = re.match(r'^(\d+)\.\t', raw)
            if m:
                n = int(m.group(1))
                if 1 <= n <= 35:
                    articles[cur] = max(articles[cur], n)

    return articles


# ── Section splitters ────────────────────────────────────────────────────────

def split_kbkf(path):
    """Split KBKF docx into main contract and ［業務委託条項］ section."""
    main, itaku, mode = [], [], 'main'
    for para in Document(path).paragraphs:
        s, t = para.style.name, para.text.strip()
        if t.startswith('［') and '業務委託条項' in t:
            mode = 'itaku'
            continue
        (itaku if mode == 'itaku' else main).append((s, t))
    return main, itaku


def split_kfkm(path):
    """Split KFKM docx into main contract and 機密保持契約書 section.
    誓約書 section (after 機密保持契約書 第14条) is omitted — each item is a top-level
    entry with no clauses, so no clause-count audit is needed.
    """
    main, nda, mode = [], [], 'main'
    for para in Document(path).paragraphs:
        s, t = para.style.name, para.text.strip()
        if t == '機密保持契約書' and mode == 'main':
            mode = 'nda'
            continue
        if '私は､当社が受託した業務' in t and mode == 'nda':
            break   # start of 誓約書 — stop here
        if mode == 'main':
            main.append((s, t))
        else:
            nda.append((s, t))
    return main, nda


# ── Comparison ───────────────────────────────────────────────────────────────

def compare(label, docx_arts, json_section, prefix=''):
    """Compare docx article structure against a JSON section.

    Args:
        docx_arts:    {第N条: clause_count} from docx parser
        json_section: sub-dict of JSON {key: {title, content, clauses}}
        prefix:       prefix to prepend when looking up JSON keys (e.g. '別紙-')
    """
    issues = []
    json_keys = {k: v for k, v in json_section.items()
                 if not k.startswith('前文') and not k.startswith('添付')}

    for jkey in json_keys:
        dkey = jkey[len(prefix):] if prefix and jkey.startswith(prefix) else jkey
        if dkey not in docx_arts:
            issues.append(f'  ❓ {jkey}: JSONにあるがdocxに見当たらない')

    for dkey in docx_arts:
        jkey = prefix + dkey if prefix else dkey
        if jkey not in json_section:
            issues.append(f'  ⚠️  {jkey}: docxにあるがJSONに無い')

    for dkey, dcnt in docx_arts.items():
        jkey = prefix + dkey if prefix else dkey
        if jkey not in json_section:
            continue
        jcnt = len(json_section[jkey].get('clauses', {}))
        if dcnt == 1 and jcnt == 0:
            continue   # single-paragraph article stored in content field — OK
        if dcnt != jcnt:
            title = json_section[jkey].get('title', '')
            issues.append(f'  ❌ {jkey} [{title}]: docx={dcnt}項  JSON={jcnt}項')

    print(f'\n  [{label}]')
    if issues:
        for i in issues:
            print(i)
    else:
        print('  ✅ 不一致なし')


def load_json(fname):
    with open(DATA / fname) as f:
        return json.load(f)


def header(docx, jf, key):
    print(f'\n{"="*70}')
    print(f'  {docx}')
    print(f'  → {jf} / {key}')
    print('=' * 70)


# ── Audit entry points ───────────────────────────────────────────────────────

def audit_bp(docx_name, json_name, contract_key):
    jd = load_json(json_name)[contract_key]
    paras = [(p.style.name, p.text.strip())
             for p in Document(DATA / docx_name).paragraphs]
    header(docx_name, json_name, contract_key)
    compare('main', extract_articles_normal(paras), jd)


def audit_bb(docx_name, json_name, contract_key):
    jd = load_json(json_name)[contract_key]
    arts = extract_articles_bb(DATA / docx_name)
    t2k = {(v.get('title') or '').strip(): k for k, v in jd.items()}

    header(docx_name, json_name, contract_key)
    issues = []
    for title, dcnt in arts.items():
        jkey = t2k.get(title)
        if not jkey:
            continue
        jcnt = len(jd[jkey].get('clauses', {}))
        if dcnt == 1 and jcnt == 0:
            continue
        if dcnt != jcnt:
            issues.append(f'  ❌ {jkey} [{title}]: docx={dcnt}項  JSON={jcnt}項')
    print('\n  [main]')
    if issues:
        for i in issues:
            print(i)
    else:
        print('  ✅ 不一致なし')


def audit_kbkf(docx_name, json_name, contract_key):
    jd = load_json(json_name)[contract_key]
    main_p, itaku_p = split_kbkf(DATA / docx_name)
    header(docx_name, json_name, contract_key)
    compare('main (第1-27条)',
            extract_articles_normal(main_p),
            {k: v for k, v in jd.items() if not k.startswith('別紙')})
    compare('別紙/業務委託条項 (第1-32条)',
            extract_articles_normal(itaku_p),
            {k: v for k, v in jd.items() if k.startswith('別紙')},
            prefix='別紙-')


def audit_kfkm(docx_name, json_name, contract_key):
    jd = load_json(json_name)[contract_key]
    main_p, nda_p = split_kfkm(DATA / docx_name)
    header(docx_name, json_name, contract_key)
    compare('main (第1-43条)',
            extract_articles_normal(main_p),
            {k: v for k, v in jd.items() if not k.startswith('添付')})
    compare('添付1 機密保持契約書 (第1-14条)',
            extract_articles_normal(nda_p),
            {k: v for k, v in jd.items() if k.startswith('添付1')},
            prefix='添付1 機密保持契約書 ')


# ── Main ─────────────────────────────────────────────────────────────────────

if __name__ == '__main__':
    audit_bp(
        'BP-IGUAZUビジネスパートナー契約書.docx',
        'customer.json',
        'IGUAZUビジネスパートナー契約書（BP）',
    )
    audit_bb(
        'BB-売買取引基本契約書.docx',
        'customer.json',
        '売買取引基本契約書（BB）',
    )
    audit_kbkf(
        'KBKF-購買基本契約書.docx',
        'supplier.json',
        '購買基本契約書(業務委託条項含む)（KB/KF）',
    )
    audit_kfkm(
        'KFKM-業務委託基本契約書.docx',
        'supplier.json',
        '業務委託基本契約書（KF/KM）',
    )
    print()
